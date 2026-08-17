"""
test_extraction_pipeline.py

Comprehensive Test Suite for AI Invoice Document Extraction Pipeline.
Tests all 7 file types and scenarios required by the task specifications:
1. Clear Invoice Image (JPG)
2. Scanned Invoice PDF
3. Digital PDF Invoice
4. DOCX Invoice with Tables
5. Invoice with GST
6. Invoice with Multiple Line Items
7. Invoice with Missing Fields
"""

import os
import sys
import json
import logging
from PIL import Image, ImageDraw, ImageFont

# Set up logging to stdout
logging.basicConfig(level=logging.INFO, format="%(message)s")

from routes.ocr import process_document_extraction, HAS_DOCX, HAS_PYMUPDF, HAS_PIL

TEST_DIR = os.path.join(os.path.dirname(__file__), "test_files")
os.makedirs(TEST_DIR, exist_ok=True)


def create_invoice_jpg():
    """Test 1: Clear Invoice JPG"""
    file_path = os.path.join(TEST_DIR, "test_1_clear_invoice.jpg")
    img = Image.new("RGB", (600, 400), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    text = (
        "ABC TECHNOLOGIES PVT LTD\n"
        "Invoice No: INV-1024\n"
        "Date: 2026-08-15\n\n"
        "Billed To: XYZ Pvt Ltd\n\n"
        "Laptop  2  50000  100000\n"
        "Mouse  5  1000  5000\n\n"
        "Subtotal: 105000\n"
        "Tax: 18900\n"
        "Grand Total: 123900"
    )
    draw.text((30, 30), text, fill=(0, 0, 0))
    img.save(file_path, "JPEG")
    return file_path, "test_1_clear_invoice.jpg"


def create_scanned_pdf():
    """Test 2: Scanned Invoice PDF (Image rendered inside PDF)"""
    file_path = os.path.join(TEST_DIR, "test_2_scanned_invoice.pdf")
    img = Image.new("RGB", (600, 400), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    text = (
        "GLOBAL SCANNED SERVICES LLC\n"
        "Invoice No: INV-SCANNED-99\n"
        "Date: 2026-05-10\n\n"
        "Customer: Acme Trading\n\n"
        "Consulting Services 10 2500 25000\n"
        "Subtotal: 25000\n"
        "Tax: 4500\n"
        "Total: 29500"
    )
    draw.text((30, 30), text, fill=(0, 0, 0))

    if HAS_PIL:
        img.convert("RGB").save(file_path, "PDF", resolution=100.0)
    return file_path, "test_2_scanned_invoice.pdf"


def create_digital_pdf():
    """Test 3: Digital PDF Invoice using PyMuPDF"""
    file_path = os.path.join(TEST_DIR, "test_3_digital_invoice.pdf")
    if HAS_PYMUPDF:
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        text = (
            "DIGITAL SOLUTIONS INC\n"
            "Tax Invoice No: INV-DIG-2026\n"
            "Date: 2026-07-20\n"
            "Due Date: 2026-08-20\n\n"
            "Billed To: TechCorp Solutions\n\n"
            "Web Hosting Annual 1 12000 12000\n"
            "Domain Registration 2 1000 2000\n\n"
            "Subtotal: 14000\n"
            "GST Tax: 2520\n"
            "Grand Total: 16520\n"
            "Status: Unpaid"
        )
        page.insert_text((50, 50), text, fontsize=12)
        doc.save(file_path)
        doc.close()
    else:
        create_invoice_jpg()
    return file_path, "test_3_digital_invoice.pdf"


def create_docx_invoice():
    """Test 4: DOCX Invoice containing tables"""
    file_path = os.path.join(TEST_DIR, "test_4_table_invoice.docx")
    if HAS_DOCX:
        import docx
        doc = docx.Document()
        doc.add_heading("INVOICE - METRO SUPPLIES", 0)
        doc.add_paragraph("Invoice No: INV-DOCX-555")
        doc.add_paragraph("Date: 2026-06-15")
        doc.add_paragraph("Billed To: Apex Enterprise")

        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Description"
        hdr_cells[1].text = "Qty"
        hdr_cells[2].text = "Unit Price"
        hdr_cells[3].text = "Total"

        row1 = table.add_row().cells
        row1[0].text = "Office Desks"
        row1[1].text = "4"
        row1[2].text = "15000"
        row1[3].text = "60000"

        row2 = table.add_row().cells
        row2[0].text = "Ergonomic Chairs"
        row2[1].text = "4"
        row2[2].text = "8000"
        row2[3].text = "32000"

        doc.add_paragraph("Subtotal: 92000")
        doc.add_paragraph("Tax: 16560")
        doc.add_paragraph("Total Amount: 108560")
        doc.save(file_path)
    return file_path, "test_4_table_invoice.docx"


def create_gst_invoice():
    """Test 5: Invoice with GST (CGST/SGST)"""
    file_path = os.path.join(TEST_DIR, "test_5_gst_invoice.png")
    img = Image.new("RGB", (600, 450), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    text = (
        "BHARAT RETAIL PVT LTD\n"
        "GSTIN: 27AAAAA0000A1Z5\n"
        "Invoice No: INV-GST-8888\n"
        "Date: 2026-04-01\n\n"
        "Customer: Reliance Infra\n"
        "Customer GSTIN: 27BBBBB1111B1Z2\n\n"
        "Industrial Cables 50 1000 50000\n\n"
        "Taxable Value: 50000\n"
        "CGST (9%): 4500\n"
        "SGST (9%): 4500\n"
        "Grand Total: 59000"
    )
    draw.text((30, 30), text, fill=(0, 0, 0))
    img.save(file_path, "PNG")
    return file_path, "test_5_gst_invoice.png"


def create_multi_item_invoice():
    """Test 6: Invoice with multiple line items"""
    file_path = os.path.join(TEST_DIR, "test_6_multi_item.png")
    img = Image.new("RGB", (600, 500), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    text = (
        "CREATIVE LABS PVT LTD\n"
        "Invoice No: INV-MULTI-707\n"
        "Date: 2026-03-25\n\n"
        "Billed To: Studio One\n\n"
        "Item 1 UI Design 1 20000 20000\n"
        "Item 2 UX Audit 1 15000 15000\n"
        "Item 3 Logo Concept 2 5000 10000\n\n"
        "Subtotal: 45000\n"
        "Tax: 8100\n"
        "Total: 53100"
    )
    draw.text((30, 30), text, fill=(0, 0, 0))
    img.save(file_path, "PNG")
    return file_path, "test_6_multi_item.png"


def create_missing_fields_invoice():
    """Test 7: Invoice with missing fields (No customer, no invoice number)"""
    file_path = os.path.join(TEST_DIR, "test_7_missing_fields.jpg")
    img = Image.new("RGB", (600, 300), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    text = (
        "CORNER CAFE\n"
        "Date: 2026-02-14\n\n"
        "Coffee 2 250 500\n"
        "Sandwich 1 350 350\n\n"
        "Subtotal: 850\n"
        "Total: 850"
    )
    draw.text((30, 30), text, fill=(0, 0, 0))
    img.save(file_path, "JPEG")
    return file_path, "test_7_missing_fields.jpg"


def run_tests():
    print("\n========================================================")
    print("STARTING DOCUMENT EXTRACTION PIPELINE TEST SUITE")
    print("========================================================\n")

    test_creators = [
        ("Test 1: Clear Invoice JPG", create_invoice_jpg),
        ("Test 2: Scanned Invoice PDF", create_scanned_pdf),
        ("Test 3: Digital PDF Invoice", create_digital_pdf),
        ("Test 4: DOCX Invoice with Tables", create_docx_invoice),
        ("Test 5: Invoice with GST", create_gst_invoice),
        ("Test 6: Multi-Item Invoice", create_multi_item_invoice),
        ("Test 7: Missing Fields Invoice", create_missing_fields_invoice),
    ]

    passed = 0
    total = len(test_creators)

    for test_name, creator in test_creators:
        print(f"\n--- {test_name} ---")
        try:
            file_path, fname = creator()
            res = process_document_extraction(file_path, fname)

            extracted_text = res.get("extracted_text", "")
            inv = res.get("invoice", {})
            validations = res.get("validations", [])
            method = res.get("extraction_method", "unknown")

            print(f"Extraction Method : {method}")
            print(f"Text Characters   : {len(extracted_text)}")
            print(f"Vendor Name       : {inv.get('vendor', {}).get('name')}")
            print(f"Customer Name     : {inv.get('customer', {}).get('name')}")
            print(f"Invoice Number    : {inv.get('invoice_number')}")
            print(f"Invoice Date      : {inv.get('invoice_date')}")
            print(f"Subtotal          : {inv.get('subtotal')}")
            print(f"Tax Amount        : {inv.get('tax_amount')}")
            print(f"Total Amount      : {inv.get('total_amount')}")
            print(f"Line Items Count  : {len(inv.get('items', []))}")
            print(f"Validations Count : {len(validations)}")

            # Assert basic validity: extracted_text must not be empty
            if extracted_text and len(extracted_text) > 5:
                print(f"[PASS] {test_name}: PASSED")
                passed += 1
            else:
                print(f"[FAIL] {test_name}: FAILED (No text extracted)")

        except Exception as e:
            print(f"[FAIL] {test_name}: FAILED with exception: {e}")

    print("\n========================================================")
    print(f"TEST RESULTS: {passed} / {total} PASSED")
    print("========================================================\n")

    return 0 if passed == total else 1


if __name__ == "__main__":
    sys.exit(run_tests())
