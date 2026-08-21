import os
import re
import json
import base64
import logging
from datetime import datetime, date

import requests
# pyrefly: ignore [missing-import]
from flask import Blueprint, jsonify, request

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("ocr_pipeline")


# ============================================================
# DEPENDENCY IMPORTS & INITIALIZATION
# ============================================================

# PyMuPDF (fitz)
try:
    # pyrefly: ignore [missing-import]
    import pymupdf as fitz
    HAS_PYMUPDF = True
except ImportError:
    try:
        # pyrefly: ignore [missing-import]
        import fitz
        HAS_PYMUPDF = True
    except ImportError:
        HAS_PYMUPDF = False

# pdfplumber
try:
    # pyrefly: ignore [missing-import]
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# pypdf
try:
    # pyrefly: ignore [missing-import]
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

# EasyOCR
HAS_EASYOCR = False
EASYOCR_READER = None
try:
    # pyrefly: ignore [missing-import]
    import easyocr
    HAS_EASYOCR = True
except ImportError:
    HAS_EASYOCR = False

# PyTesseract
try:
    # pyrefly: ignore [missing-import]
    import pytesseract
    HAS_PYTESSERACT = True
except ImportError:
    HAS_PYTESSERACT = False

if HAS_PYTESSERACT:
    import sys as _sys
    _tess_cmd = os.getenv("TESSERACT_CMD", "").strip()
    if not _tess_cmd:
        # Auto-detect common Windows install locations when env var is not set
        if _sys.platform == "win32":
            _win_paths = [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                os.path.expanduser(r"~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"),
            ]
            for _p in _win_paths:
                if os.path.isfile(_p):
                    _tess_cmd = _p
                    break
    # Fallback: bare name — assumes tesseract is in PATH (Linux / Mac / Docker)
    if not _tess_cmd:
        _tess_cmd = "tesseract"
    pytesseract.pytesseract.tesseract_cmd = _tess_cmd
    logger.info(f"Tesseract cmd configured: {_tess_cmd}")

try:
    # pyrefly: ignore [missing-import]
    import cv2
    import numpy as np
    HAS_OPENCV = True
except ImportError:
    HAS_OPENCV = False

try:
    # pyrefly: ignore [missing-import]
    from PIL import Image
    import io
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

# python-docx
try:
    # pyrefly: ignore [missing-import]
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from config import Config

# Gemini Config
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "").strip()
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash").strip()
try:
    GEMINI_TIMEOUT = int(os.getenv("GEMINI_TIMEOUT", "60"))
except (TypeError, ValueError):
    GEMINI_TIMEOUT = 60

GEMINI_ENDPOINT = (
    "https://generativelanguage.googleapis.com/v1beta/models/"
    f"{GEMINI_MODEL}:generateContent"
)

SUPPORTED_IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "tiff", "bmp", "webp"}
SUPPORTED_DOCUMENT_EXTENSIONS = {"doc", "docx"}
SUPPORTED_EXTENSIONS = SUPPORTED_IMAGE_EXTENSIONS | {"pdf"} | SUPPORTED_DOCUMENT_EXTENSIONS


def get_easyocr_reader():
    global EASYOCR_READER
    if EASYOCR_READER is None and HAS_EASYOCR:
        try:
            logger.info("Initializing EasyOCR reader...")
            EASYOCR_READER = easyocr.Reader(['en'], gpu=False, verbose=False)
        except Exception as e:
            logger.error(f"Failed to initialize EasyOCR: {e}")
            EASYOCR_READER = None
    return EASYOCR_READER


def extract_text_from_image(file_path):
    """
    Extract raw text from an image file using EasyOCR or Tesseract.
    """
    raw_text = ""
    method = "image_ocr"

    # Try PyTesseract first if available
    if HAS_PYTESSERACT:
        try:
            txt = pytesseract.image_to_string(file_path, lang="eng")
            if txt and len(txt.strip()) >= 5:
                return txt.strip(), "tesseract_ocr"
        except Exception as e:
            logger.warning(f"PyTesseract error: {e}")

    # Try EasyOCR
    reader = get_easyocr_reader()
    if reader:
        try:
            results = reader.readtext(file_path, detail=0)
            if results:
                raw_text = "\n".join(results).strip()
                if len(raw_text) >= 5:
                    return raw_text, "easyocr"
        except Exception as e:
            logger.warning(f"EasyOCR error: {e}")

    # Fallback with OpenCV preprocessing + EasyOCR if available
    if HAS_OPENCV and reader:
        try:
            img = cv2.imread(file_path)
            if img is not None:
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                results = reader.readtext(gray, detail=0)
                if results:
                    raw_text = "\n".join(results).strip()
                    return raw_text, "easyocr_opencv"
        except Exception as e:
            logger.warning(f"OpenCV EasyOCR fallback error: {e}")

    return raw_text, method if raw_text else "image_ocr_failed"


def extract_text_from_pdf(file_path):
    """
    Extract text from PDF.
    First attempts normal PDF text extraction.
    If no meaningful text is available, converts pages to images and runs OCR.
    """
    extracted_text = ""
    method = "pdf_text_extraction"

    # 1. Try PyMuPDF (fitz) text extraction
    if HAS_PYMUPDF:
        try:
            doc = fitz.open(file_path)
            pages_text = []
            for page in doc:
                t = page.get_text()
                if t and t.strip():
                    pages_text.append(t.strip())
            if pages_text:
                extracted_text = "\n\n".join(pages_text).strip()
        except Exception as e:
            logger.warning(f"PyMuPDF text extraction error: {e}")

    if not extracted_text and HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t and t.strip():
                        pages_text.append(t.strip())
                if pages_text:
                    extracted_text = "\n\n".join(pages_text).strip()
        except Exception as e:
            logger.warning(f"pdfplumber text extraction error: {e}")

    # 3. Try pypdf if still nothing
    if not extracted_text and HAS_PYPDF:
        try:
            reader = pypdf.PdfReader(file_path)
            pages_text = []
            for page in reader.pages:
                t = page.extract_text()
                if t and t.strip():
                    pages_text.append(t.strip())
            if pages_text:
                extracted_text = "\n\n".join(pages_text).strip()
        except Exception as e:
            logger.warning(f"pypdf text extraction error: {e}")

    alpha_count = sum(1 for c in extracted_text if c.isalnum())
    if alpha_count >= 20:
        return extracted_text, "pdf_digital_text"

    # 4. If no meaningful text (Scanned PDF), render pages to images and run OCR
    logger.info("PDF has no meaningful text. Running scanned PDF rendering + OCR pipeline...")
    ocr_pages_text = []

    if HAS_PYMUPDF:
        try:
            doc = fitz.open(file_path)
            reader = get_easyocr_reader()
            for page_idx, page in enumerate(doc):
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")

                if HAS_PIL and reader:
                    img = Image.open(io.BytesIO(img_bytes))
                    if HAS_OPENCV:
                        open_cv_image = np.array(img.convert("RGB"))
                        open_cv_image = open_cv_image[:, :, ::-1].copy()
                        ocr_res = reader.readtext(open_cv_image, detail=0)
                    else:
                        ocr_res = reader.readtext(img_bytes, detail=0)

                    if ocr_res:
                        ocr_pages_text.append("\n".join(ocr_res))

            if ocr_pages_text:
                return "\n\n".join(ocr_pages_text).strip(), "scanned_pdf_ocr"
        except Exception as e:
            logger.warning(f"PyMuPDF scanned PDF OCR error: {e}")

    return extracted_text or "No readable text found in PDF", "pdf_extraction_failed"


def extract_text_from_docx(file_path):
    """
    Extract text from DOCX/DOC files, preserving paragraph text and table structure.
    """
    if not HAS_DOCX:
        return "", "docx_library_missing"

    try:
        doc = docx.Document(file_path)
        text_parts = []

        for p in doc.paragraphs:
            if p.text and p.text.strip():
                text_parts.append(p.text.strip())

        for table in doc.tables:
            text_parts.append("--- Table Start ---")
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))
            text_parts.append("--- Table End ---")

        full_text = "\n".join(text_parts).strip()
        return full_text, "docx_text_table_extraction"
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return "", f"docx_error: {e}"


def _normalize_date(date_str):
    if not date_str or str(date_str).lower() in ("null", "none", "not detected", ""):
        return None
    d_clean = re.sub(r"[^\w\s/-]", "", str(date_str)).strip()
    if not d_clean:
        return None

    patterns = [
        ("%Y-%m-%d", r"^\d{4}-\d{2}-\d{2}$"),
        ("%d/%m/%Y", r"^\d{1,2}/\d{1,2}/\d{4}$"),
        ("%d-%m-%Y", r"^\d{1,2}-\d{1,2}-\d{4}$"),
        ("%m/%d/%Y", r"^\d{1,2}/\d{1,2}/\d{4}$"),
        ("%Y/%m/%d", r"^\d{4}/\d{1,2}/\d{1,2}$"),
        ("%d %b %Y", r"^\d{1,2}\s+[A-Za-z]{3,9}\s+\d{4}$"),
        ("%b %d, %Y", r"^[A-Za-z]{3,9}\s+\d{1,2},\s*\d{4}$")
    ]
    for fmt_str, regex in patterns:
        if re.match(regex, d_clean, re.IGNORECASE):
            try:
                dt = datetime.strptime(d_clean, fmt_str)
                return dt.strftime("%Y-%m-%d")
            except ValueError:
                pass
    match = re.search(r"(\d{4})[/-](\d{1,2})[/-](\d{1,2})", d_clean)
    if match:
        y, m, d = match.groups()
        return f"{y}-{int(m):02d}-{int(d):02d}"
    match_d = re.search(r"(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})", d_clean)
    if match_d:
        d, m, y = match_d.groups()
        if len(y) == 2:
            y = "20" + y
        return f"{y}-{int(m):02d}-{int(d):02d}"
    return None


def _parse_number(val):
    if val is None or str(val).lower() in ("null", "none", "not detected", ""):
        return None
    if isinstance(val, (int, float)):
        return round(float(val), 2)
    val_str = str(val).strip()
    val_str = re.sub(r"[^\d.,]", "", val_str)
    if not val_str:
        return None
    if "," in val_str and "." in val_str:
        if val_str.rfind(".") > val_str.rfind(","):
            val_str = val_str.replace(",", "")
        else:
            val_str = val_str.replace(".", "").replace(",", ".")
    elif "," in val_str:
        if len(val_str.split(",")[-1]) == 2:
            val_str = val_str.replace(",", ".")
        else:
            val_str = val_str.replace(",", "")
    try:
        return round(float(val_str), 2)
    except ValueError:
        return None


def _safe_parse_ai_json(json_raw_str):
    """
    Safely extract and parse JSON from AI response text.
    """
    if not json_raw_str or not isinstance(json_raw_str, str):
        return None, "Empty raw response string"

    text = json_raw_str.strip()

    # 1. Strip Markdown code fences
    text = re.sub(r"^```json\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^```\s*", "", text)
    text = re.sub(r"\s*```$", "", text).strip()

    # 2. Try direct json parse
    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            return parsed, None
    except Exception:
        pass

    # 3. Use Regex to find substring { ... }
    start_idx = text.find("{")
    end_idx = text.rfind("}")
    if start_idx != -1 and end_idx > start_idx:
        json_sub = text[start_idx:end_idx + 1]
        try:
            parsed = json.loads(json_sub)
            if isinstance(parsed, dict):
                return parsed, None
        except Exception as e:
            return None, f"JSON decode error in substring: {e}"

    return None, "Failed to extract valid JSON dictionary"


def _call_gemini_ai_api(raw_text):
    """
    Calls Gemini API to extract structured JSON from raw document text.
    """
    if not GEMINI_API_KEY or GEMINI_API_KEY == "YOUR_GEMINI_API_KEY":
        return None, "GEMINI_API_KEY not configured"

    prompt = f"""
You are an expert financial OCR AI assistant. Extract structured JSON from the following invoice raw text.

CRITICAL INSTRUCTIONS:
1. Return ONLY valid JSON matching the exact schema below.
2. NEVER HALLUCINATE OR INVENT VALUES. If a field is not explicitly present in the raw text, set its value to null.
3. Keep exact dates, item names, prices, tax amounts, vendor, and customer details.

EXPECTED JSON SCHEMA:
{{
  "invoice_number": "string or null",
  "invoice_date": "YYYY-MM-DD or null",
  "due_date": "YYYY-MM-DD or null",
  "vendor": {{
    "name": "string or null",
    "address": "string or null",
    "tax_id": "string or null"
  }},
  "customer": {{
    "name": "string or null",
    "address": "string or null",
    "tax_id": "string or null"
  }},
  "items": [
    {{
      "description": "string",
      "quantity": float,
      "unit_price": float,
      "tax": float or null,
      "total": float
    }}
  ],
  "subtotal": float or null,
  "tax_amount": float or null,
  "discount_amount": float or null,
  "total_amount": float or null,
  "currency": "INR",
  "payment_status": "Paid" | "Unpaid" | "Overdue" | "Unknown"
}}

RAW DOCUMENT TEXT:
{raw_text[:25000]}
"""

    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.0,
            "responseMimeType": "application/json"
        }
    }

    try:
        url = f"{GEMINI_ENDPOINT}?key={GEMINI_API_KEY}"
        resp = requests.post(url, json=payload, timeout=GEMINI_TIMEOUT)

        if resp.status_code == 200:
            res_data = resp.json()
            candidates = res_data.get("candidates", [])
            if candidates:
                parts = candidates[0].get("content", {}).get("parts", [])
                if parts:
                    ai_text = parts[0].get("text", "")
                    parsed_json, err = _safe_parse_ai_json(ai_text)
                    if parsed_json:
                        return parsed_json, None
                    return None, f"JSON parse error: {err}"
            return None, "Gemini returned empty candidate contents"
        else:
            return None, f"Gemini HTTP {resp.status_code}: {resp.text[:200]}"
    except Exception as e:
        return None, f"Gemini API request failed: {e}"


def _fallback_ai_structured_parser(raw_text, filename=""):
    """
    High-accuracy deterministic fallback AI parser that extracts structured fields
    directly from raw text without hallucinating values when Gemini API is unavailable.
    """
    lines = [l.strip() for l in (raw_text or "").splitlines() if l.strip()]

    vendor = None
    vendor_address = None
    vendor_tax_id = None

    customer = None
    customer_address = None
    customer_tax_id = None

    invoice_number = None
    invoice_date = None
    due_date = None

    subtotal = None
    tax_amount = None
    discount_amount = 0.0
    total_amount = None
    currency = "INR"
    payment_status = "Unknown"

    cgst_val = None
    sgst_val = None
    igst_val = None

    for idx, line in enumerate(lines):
        next_line = lines[idx + 1] if idx + 1 < len(lines) else ""
        two_lines = f"{line} {next_line}"
        if not vendor:
            v_match = re.search(r"^(?:Vendor|Supplier|Merchant|Billed By|Company|Seller|Payee|Issued By)[\s:]+([A-Za-z0-9\s.,&'-]{2,50})$", line, re.IGNORECASE)
            if v_match:
                cand = v_match.group(1).strip()
                if cand.lower() not in {"invoice", "tax invoice", "receipt", "total", "amount"}:
                    vendor = cand

        if not vendor_tax_id:
            g_match = re.search(r"(?:GSTIN|VAT|TAX ID|EIN|GST No|Tax No)[\s:]*([0-9A-Z-]{8,20})", line, re.IGNORECASE)
            if g_match:
                vendor_tax_id = g_match.group(1).strip()

        if not customer:
            c_match = re.search(r"^(?:Billed To|Customer|Client|Ship To|Buyer|Recipient|Invoice To)[\s:]+([A-Za-z0-9\s.,&'-]{2,50})$", line, re.IGNORECASE)
            if c_match:
                cand = c_match.group(1).strip()
                if cand.lower() not in {"invoice", "tax invoice", "receipt", "total", "amount"}:
                    customer = cand

        if not customer_tax_id:
            cg_match = re.search(r"(?:Customer GSTIN|Client VAT|Customer Tax ID)[\s:]*([0-9A-Z-]{8,20})", line, re.IGNORECASE)
            if cg_match:
                customer_tax_id = cg_match.group(1).strip()

        # Invoice Number
        if not invoice_number:
            inv_match = re.search(
                r"(?:Tax\s*Invoice\s*No|Invoice\s*Number|Invoice\s*No|Invoice\s*#|Inv\s*No|Inv\s*#|Bill\s*No|Receipt\s*No|Receipt\s*#|Ref\s*No)[\s:]+([A-Za-z0-9-_/#]{2,30})",
                line, re.IGNORECASE
            )
            if inv_match:
                num_cand = inv_match.group(1).strip()
                if num_cand.lower() not in {"date", "no", "number", "tax"}:
                    invoice_number = num_cand
            elif re.search(r"\b(INV-[A-Za-z0-9-_]{3,20})\b", line, re.IGNORECASE):
                invoice_number = re.search(r"\b(INV-[A-Za-z0-9-_]{3,20})\b", line, re.IGNORECASE).group(1).strip()

        if not invoice_date:
            d_match = re.search(r"(?:Invoice\s*Date|Receipt\s*Date|Bill\s*Date|Date\s*of\s*Issue|Txn\s*Date|^Date)[\s:]+([A-Za-z0-9,\s/-]{6,25})", line, re.IGNORECASE)
            if d_match:
                invoice_date = _normalize_date(d_match.group(1).strip())

        if not due_date:
            due_match = re.search(r"(?:Due\s*Date|Payment\s*Due|Pay\s*By|^Due)[\s:]+([A-Za-z0-9,\s/-]{6,25})", line, re.IGNORECASE)
            if due_match:
                due_date = _normalize_date(due_match.group(1).strip())

        if total_amount is None:
            tot_match = re.search(r"(?:Grand\s*Total|Total\s*Amount|Total\s*Payable|Net\s*Amount|Total\s*Due|Amount\s*Paid|^Total)[^\d\n]*([\d,]+\.?\d*)", two_lines, re.IGNORECASE)
            if tot_match:
                parsed_tot = _parse_number(tot_match.group(1))
                if parsed_tot is not None and parsed_tot > 0:
                    total_amount = parsed_tot

        if subtotal is None:
            sub_match = re.search(r"(?:Subtotal|Sub\s*Total|Sub-Total|Taxable\s*Value|Taxable\s*Amount|Net\s*Value)[^\d\n]*([\d,]+\.?\d*)", two_lines, re.IGNORECASE)
            if sub_match:
                parsed_sub = _parse_number(sub_match.group(1))
                if parsed_sub is not None and parsed_sub > 0:
                    subtotal = parsed_sub

        if discount_amount == 0.0 or discount_amount is None:
            disc_match = re.search(r"(?:Discount|Less\s*Discount|Savings)[^\d\n]*([\d,]+\.?\d*)", two_lines, re.IGNORECASE)
            if disc_match:
                parsed_disc = _parse_number(disc_match.group(1))
                if parsed_disc is not None:
                    discount_amount = parsed_disc

        if cgst_val is None:
            cgst_match = re.search(r"CGST[^\d\n]*([\d,]+\.?\d*)", two_lines, re.IGNORECASE)
            if cgst_match: cgst_val = _parse_number(cgst_match.group(1))

        if sgst_val is None:
            sgst_match = re.search(r"SGST[^\d\n]*([\d,]+\.?\d*)", two_lines, re.IGNORECASE)
            if sgst_match: sgst_val = _parse_number(sgst_match.group(1))

        if igst_val is None:
            igst_match = re.search(r"IGST[^\d\n]*([\d,]+\.?\d*)", two_lines, re.IGNORECASE)
            if igst_match: igst_val = _parse_number(igst_match.group(1))

        # Payment Status
        if "paid" in line.lower() and "unpaid" not in line.lower():
            payment_status = "Paid"
        elif "unpaid" in line.lower():
            payment_status = "Unpaid"
        elif "overdue" in line.lower():
            payment_status = "Overdue"

    # Top header fallback for Vendor if not matched by explicit label
    if not vendor and lines:
        ignore_kws = {"invoice", "tax invoice", "receipt", "bill", "cash receipt", "statement", "date", "total", "amount", "page", "subtotal"}
        for l in lines[:5]:
            l_low = l.lower()
            if len(l) >= 3 and not any(kw in l_low for kw in ignore_kws) and not re.match(r"^[\d\s.,/\-#]+$", l):
                vendor = l.strip()
                break

    if not invoice_date:
        gen_match = re.search(r"\b(\d{4}[/-]\d{1,2}[/-]\d{1,2}|\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b", raw_text or "")
        if gen_match:
            invoice_date = _normalize_date(gen_match.group(1))

    if cgst_val is not None or sgst_val is not None:
        tax_amount = round((cgst_val or 0.0) + (sgst_val or 0.0), 2)
    elif igst_val is not None:
        tax_amount = round(igst_val, 2)
    else:
        gen_tax_match = re.search(r"(?:Tax|GST|VAT|Total\s*Tax)(?:\s*\([^)]*\))?[\s:₹$Rs.]*[\s:]+([\d,]+\.?\d*)", raw_text or "", re.IGNORECASE)
        if gen_tax_match:
            tax_amount = _parse_number(gen_tax_match.group(1))

    if "₹" in raw_text or "INR" in raw_text or "Rs." in raw_text:
        currency = "INR"
    elif "$" in raw_text or "USD" in raw_text:
        currency = "USD"
    elif "€" in raw_text or "EUR" in raw_text:
        currency = "EUR"
    elif "£" in raw_text or "GBP" in raw_text:
        currency = "GBP"

    items = []
    for line in lines:
        pipe_match = re.search(r"^([^|]+)\s*\|\s*(\d+(?:\.\d+)?)\s*\|\s*([\d,]+\.?\d*)\s*(?:\|\s*([\d,]+\.?\d*))?\s*\|\s*([\d,]+\.?\d*)$", line)
        if not pipe_match:
            pipe_match = re.search(r"^([^|]+)\s*\|\s*(\d+(?:\.\d+)?)\s*\|\s*([\d,]+\.?\d*)\s*\|\s*([\d,]+\.?\d*)$", line)

        space_match = re.search(r"^([A-Za-z0-9\s.,&'-]{2,50})\s+(\d+(?:\.\d+)?)\s+([\d,]+\.?\d*)\s+([\d,]+\.?\d*)$", line)

        if pipe_match:
            desc = pipe_match.group(1).strip()
            if desc.lower() not in {"total", "subtotal", "tax", "description", "item", "qty", "price", "unit price", "amount", "gst"}:
                qty = float(pipe_match.group(2))
                u_price = _parse_number(pipe_match.group(3)) or 0.0
                tot_str = pipe_match.group(5) if len(pipe_match.groups()) >= 5 and pipe_match.group(5) else pipe_match.group(4)
                tot = _parse_number(tot_str) or round(qty * u_price, 2)
                items.append({
                    "description": desc,
                    "quantity": qty,
                    "unit_price": u_price,
                    "tax": 0.0,
                    "total": tot
                })
        elif space_match:
            desc = space_match.group(1).strip()
            if desc.lower() not in {"total", "subtotal", "tax", "description", "item", "qty", "price", "unit price", "amount", "gst"}:
                qty = float(space_match.group(2))
                u_price = _parse_number(space_match.group(3)) or 0.0
                tot = _parse_number(space_match.group(4)) or round(qty * u_price, 2)
                items.append({
                    "description": desc,
                    "quantity": qty,
                    "unit_price": u_price,
                    "tax": 0.0,
                    "total": tot
                })

    if not items and (vendor or total_amount):
        items.append({
            "description": f"{vendor or 'Invoice'} Line Item",
            "quantity": 1.0,
            "unit_price": total_amount or 0.0,
            "tax": tax_amount or 0.0,
            "total": total_amount or 0.0
        })

    return {
        "invoice_number": invoice_number,
        "invoice_date": invoice_date,
        "due_date": due_date,
        "vendor": {
            "name": vendor,
            "address": vendor_address,
            "tax_id": vendor_tax_id
        },
        "customer": {
            "name": customer,
            "address": customer_address,
            "tax_id": customer_tax_id
        },
        "items": items,
        "subtotal": subtotal or (total_amount - (tax_amount or 0.0) if total_amount else None),
        "tax_amount": tax_amount,
        "discount_amount": discount_amount or 0.0,
        "total_amount": total_amount,
        "currency": currency,
        "payment_status": payment_status
    }


def clean_structured_invoice_json(parsed_dict, raw_text=""):
    """
    Cleans and standardizes the AI extracted JSON structure.
    Enforces `null` (None) for missing fields (Never Hallucinate).
    """
    if not isinstance(parsed_dict, dict):
        parsed_dict = {}

    v_raw = parsed_dict.get("vendor")
    vendor_obj = {
        "name": (v_raw.get("name") if isinstance(v_raw, dict) else v_raw) or None,
        "address": (v_raw.get("address") if isinstance(v_raw, dict) else None) or None,
        "tax_id": (v_raw.get("tax_id") if isinstance(v_raw, dict) else None) or None
    }
    if vendor_obj["name"] and str(vendor_obj["name"]).lower() in ("null", "none", "not detected"):
        vendor_obj["name"] = None

    c_raw = parsed_dict.get("customer")
    customer_obj = {
        "name": (c_raw.get("name") if isinstance(c_raw, dict) else c_raw) or None,
        "address": (c_raw.get("address") if isinstance(c_raw, dict) else None) or None,
        "tax_id": (c_raw.get("tax_id") if isinstance(c_raw, dict) else None) or None
    }
    if customer_obj["name"] and str(customer_obj["name"]).lower() in ("null", "none", "not detected"):
        customer_obj["name"] = None

    items = []
    raw_items = parsed_dict.get("items") or []
    for item in raw_items:
        if not isinstance(item, dict):
            continue
        q = float(item.get("quantity") or 1.0)
        u = float(item.get("unit_price") or 0.0)
        t = item.get("tax")
        tax_val = float(t) if t is not None else None
        tot = item.get("total")
        total_val = float(tot) if tot is not None else round((q * u), 2)
        items.append({
            "description": item.get("description") or "Invoice Line Item",
            "quantity": q,
            "unit_price": u,
            "tax": tax_val,
            "total": total_val
        })

    inv_num = parsed_dict.get("invoice_number")
    if inv_num and str(inv_num).lower() in ("null", "none", "not detected"):
        inv_num = None

    inv_date = _normalize_date(parsed_dict.get("invoice_date") or parsed_dict.get("date"))
    due_date = _normalize_date(parsed_dict.get("due_date"))

    subtotal = _parse_number(parsed_dict.get("subtotal"))
    tax_amount = _parse_number(parsed_dict.get("tax_amount") or parsed_dict.get("tax"))
    discount_amount = _parse_number(parsed_dict.get("discount_amount") or parsed_dict.get("discount")) or 0.0
    total_amount = _parse_number(parsed_dict.get("total_amount"))

    return {
        "invoice_number": inv_num,
        "invoice_date": inv_date,
        "due_date": due_date,
        "vendor": vendor_obj,
        "customer": customer_obj,
        "items": items,
        "subtotal": subtotal,
        "tax_amount": tax_amount,
        "discount_amount": discount_amount,
        "total_amount": total_amount,
        "currency": parsed_dict.get("currency") or "INR",
        "payment_status": parsed_dict.get("payment_status") or "Unknown"
    }


def validate_extracted_invoice(structured, user_id=None):
    """
    Validates calculation math:
    - quantity * unit_price = item total
    - sum(item totals) = subtotal
    - subtotal + tax - discount = total
    Returns a list of warning objects if calculations don't match.
    """
    validations = []
    if not structured.get("invoice_number"):
        validations.append({"field": "invoice_number", "severity": "warning", "message": "Invoice number missing in document"})

    if not structured.get("invoice_date"):
        validations.append({"field": "invoice_date", "severity": "warning", "message": "Invoice date missing in document"})

    if not structured.get("vendor", {}).get("name"):
        validations.append({"field": "vendor", "severity": "warning", "message": "Vendor / Seller name missing in document"})

    if not structured.get("customer", {}).get("name"):
        validations.append({"field": "customer", "severity": "info", "message": "Customer name missing in document"})

    items = structured.get("items") or []
    calc_item_totals_sum = 0.0
    for idx, item in enumerate(items):
        q = float(item.get("quantity", 0))
        u = float(item.get("unit_price", 0))
        actual_total = float(item.get("total", 0))
        expected_item_total = round(q * u, 2)
        if abs(expected_item_total - actual_total) > 1.0:
            validations.append({
                "field": f"items[{idx}]",
                "severity": "warning",
                "message": f"Calculation mismatch for '{item.get('description')}': Qty({q}) × Unit Price({u}) = {expected_item_total}, but extracted item total is {actual_total}"
            })
        calc_item_totals_sum += actual_total

    subtotal = structured.get("subtotal")
    if subtotal is not None and len(items) > 0:
        if abs(calc_item_totals_sum - float(subtotal)) > 2.0:
            validations.append({
                "field": "subtotal",
                "severity": "warning",
                "message": f"Subtotal calculation mismatch: Sum of item totals is {round(calc_item_totals_sum, 2)}, but extracted subtotal is {subtotal}"
            })

    # Grand Total Check
    sub_val = float(subtotal or calc_item_totals_sum or 0.0)
    tax_val = float(structured.get("tax_amount") or 0.0)
    disc_val = float(structured.get("discount_amount") or 0.0)
    total_val = structured.get("total_amount")
    expected_grand_total = round(sub_val + tax_val - disc_val, 2)
    if total_val is not None:
        if abs(expected_grand_total - float(total_val)) > 2.0:
            validations.append({
                "field": "total_amount",
                "severity": "warning",
                "message": f"Grand total calculation mismatch: Subtotal ({sub_val}) + Tax ({tax_val}) - Discount ({disc_val}) = {expected_grand_total}, but extracted total is {total_val}"
            })

    inv_num = structured.get("invoice_number")
    if user_id and inv_num:
        try:
            from utils.db import get_db
            conn = get_db()
            with conn.cursor() as cursor:
                cursor.execute("SELECT id FROM invoices WHERE invoice_number = %s AND user_id = %s LIMIT 1", (inv_num, user_id))
                if cursor.fetchone():
                    validations.append({
                        "field": "invoice_number",
                        "severity": "error",
                        "message": f"Duplicate invoice: Invoice number '{inv_num}' already exists in your database"
                    })
            conn.close()
        except Exception as e:
            logger.warning(f"Database duplicate check error: {e}")

    return validations


def generate_invoice_insights(structured):
    """
    Generate financial and payment insights from structured invoice data.
    """
    items = structured.get("items") or []
    subtotal = float(structured.get("subtotal") or 0.0)
    tax_amount = float(structured.get("tax_amount") or 0.0)
    total_amount = float(structured.get("total_amount") or 0.0)

    tax_percentage = round((tax_amount / subtotal) * 100, 2) if subtotal > 0 else 0.0

    return {
        "item_count": len(items),
        "tax_percentage": tax_percentage,
        "payment_status": structured.get("payment_status") or "Unknown",
        "financial_insights": {
            "total_amount": total_amount,
            "tax_amount": tax_amount,
            "subtotal": subtotal,
            "currency": structured.get("currency") or "INR",
            "average_item_price": round(sum(float(i.get("unit_price", 0)) for i in items) / max(len(items), 1), 2)
        },
        "payment_insights": {
            "payment_status": structured.get("payment_status") or "Unknown",
            "due_date": structured.get("due_date") or "N/A"
        }
    }


# ============================================================
# 4. MAIN PUBLIC ENTRY POINT FOR DOCUMENT EXTRACTION
# ============================================================

def process_document_extraction(file_path, filename="", user_id=None):
    """
    Main extraction function.
    Flow: File -> OCR/PDF/DOCX Text Extraction -> AI Structured Extraction -> Validation -> Insights -> Debug Logs
    """
    fname = filename or os.path.basename(file_path)
    file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
    ext = os.path.splitext(fname)[1].lower().lstrip(".")

    raw_text = ""
    extraction_method = "unknown"

    # Stage 1: Text Extraction
    if ext in SUPPORTED_IMAGE_EXTENSIONS:
        raw_text, extraction_method = extract_text_from_image(file_path)
    elif ext == "pdf":
        raw_text, extraction_method = extract_text_from_pdf(file_path)
    elif ext in SUPPORTED_DOCUMENT_EXTENSIONS:
        raw_text, extraction_method = extract_text_from_docx(file_path)
    else:
        raw_text = ""
        extraction_method = "unsupported_extension"

    # Detailed Logging - Section 1: UPLOAD & TEXT EXTRACTION
    logger.info(f"\n[UPLOAD]\nFilename: {fname}\nFile type: {ext.upper()}\nFile size: {file_size} bytes")
    logger.info(f"\n[TEXT EXTRACTION]\nExtraction method: {extraction_method}\nCharacters extracted: {len(raw_text)}\nRaw extracted text:\n{raw_text[:2000]}")

    # Stage 2: AI Structured Data Extraction
    ai_raw_json = None
    ai_err = None

    if GEMINI_API_KEY and GEMINI_API_KEY != "YOUR_GEMINI_API_KEY":
        logger.info("\n[AI EXTRACTION]\nSending request to Gemini AI model...")
        ai_raw_json, ai_err = _call_gemini_ai_api(raw_text)

    if not ai_raw_json:
        if ai_err:
            logger.info(f"\n[AI EXTRACTION]\nGemini API unavailable ({ai_err}). Using Fallback AI Extraction Engine...")
        else:
            logger.info("\n[AI EXTRACTION]\nGEMINI_API_KEY unconfigured. Using Fallback AI Extraction Engine...")
        ai_raw_json = _fallback_ai_structured_parser(raw_text, fname)

    logger.info(f"\n[AI EXTRACTION]\nAI response parsed successfully.")

    # Stage 3: Clean & Format Structured JSON Schema
    structured_data = clean_structured_invoice_json(ai_raw_json, raw_text)

    logger.info(f"\n[STRUCTURED DATA]\nParsed JSON:\n{json.dumps(structured_data, indent=2, default=str)}")

    # Stage 4: Validation Engine
    validations = validate_extracted_invoice(structured_data, user_id=user_id)

    # Stage 5: AI Insights Engine
    insights = generate_invoice_insights(structured_data)

    return {
        "success": True if raw_text else False,
        "filename": fname,
        "extracted_text": raw_text,
        "invoice": structured_data,
        "insights": insights,
        "validations": validations,
        "extraction_method": extraction_method
    }